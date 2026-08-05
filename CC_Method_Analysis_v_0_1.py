# Will require some dependencies
import pandas as pd
import numpy as np
from uncertainties import ufloat

from scipy import stats
from scipy.stats import shapiro
from scipy.stats import linregress
from scipy.stats import gaussian_kde # for reference interval 

import statsmodels.api as sm
from statsmodels.formula.api import ols

import pingouin as pg
from pingouin import ttest
from pingouin import wilcoxon

from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from pygam import LinearGAM, s
from statsmodels.nonparametric.smoothers_lowess import lowess


from PIL import Image

#plotly libraries
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots

# libraries for creating directory for figure
import os
import csv
import shutil


def manage_figures_folder():
    """
    This function manages the Figures folder by deleting all existing contents if the folder exists,
    or creating the folder if it does not exist.
    """
    import shutil
    from pathlib import Path

    folder = Path("Figures")
    
    if folder.exists():
        shutil.rmtree(folder)   # deletes everything including the folder itself
    folder.mkdir(exist_ok=False)  # creates it fresh


def FigWhiteCrop(img_file):
    """
    This function crops the image to the bounding box of non-white pixels and saves it with a new name.
    It saves the cropped image in the Figures directory, with file name prefixed with "cropped_plot_".
    """
    # Open the image file
    img = Image.open(img_file)

    # Convert image to numpy array
    img_array = np.array(img)

    # Find non-white pixels
    non_white_pixels = np.where(img_array[:, :, :3] != 255)

    # Get bounding box coordinates
    top_left = np.min(non_white_pixels, axis=1)
    bottom_right = np.max(non_white_pixels, axis=1)

    # margin in pixels to leave
    margin = 10
    # Crop the image
    cropped_img = img.crop(
            (top_left[1] - margin,
            top_left[0] - margin,
            bottom_right[1] + margin,
            bottom_right[0] + margin,
            )
         )

    # Extract filename without extension
    base_name = os.path.splitext(os.path.basename(img_file))[0]

    # Save the cropped image
    # cropped_img.save("Figures/cropped_plot.png")
    cropped_img.save(f"Figures/cropped_plot_{base_name}.png")

"""
Below is modififed version of the function above
where suffix is added to output filename so as to not overwrite the previous output.
"""

def FigWhiteCrop2(img_file):
    """
    Function to crop uncessary white space in figures.
    In contrast to FigWhiteCrop, this function FigWhiteCrop2 keeps the existing input file name as the prefix and concatenates an index value as a suffix
    """
    # Open the image file
    img = Image.open(img_file)

    # Convert image to numpy array
    img_array = np.array(img)

    # Find non-white pixels
    non_white_pixels = np.where(img_array[:, :, :3] != 255)

    # Get bounding box coordinates
    top_left = np.min(non_white_pixels, axis=1)
    bottom_right = np.max(non_white_pixels, axis=1)

    # margin in pixels to leave
    margin = 10
    # Crop the image
    cropped_img = img.crop(
        (
            top_left[1] - margin,
            top_left[0] - margin,
            bottom_right[1] + margin,
            bottom_right[0] + margin,
        )
    )

    # Extract filename without extension
    base_name, ext = os.path.splitext(os.path.basename(img_file))

    # Create the base path for the cropped image
    base_path = f"Figures/cropped_plot_{base_name}{ext}"

    # Add suffix number if the file already exists
    i = 1
    while os.path.exists(base_path):
        base_path = f"Figures/cropped_plot_{base_name}_{i}{ext}"
        i += 1
    # Save the cropped image with potential suffix
    cropped_img.save(base_path)


def sample_comparions_tableX(x, y, z, **error_info):
    """
    This function generates a table looking at the difference & % difference between the methods.
    Based on error compute if the difference is within acceptable error.

    The error can be defined as total allowable error or per IQMH terminology of allowable performance limits(APL),
    or even precision goal.

    This function is designed to take a single concentration cut-off and two error thresholds.
    One error threshold, the first error1 variable is to be DEFINED IN ABSOLUTE NUMBERS.
    The other error value, error2 is to be DEFINED IN PERCENTAGE TERMS.
    When both error threshold is provided, the concentration_cut MUST BE PROVIDED.

    When concentration cut-off is provided, the error1 term is for absolute error below the cut-off and the
    error2 term is applied above the cut-off

    x = reference method
    y = test method
    z = hue or characteristic or unique to this sample
    """
    error1 = error_info.get("error1", None)
    error2 = error_info.get("error2", None)
    Error_level_cut_off = error_info.get("Error_level_cut_off", None)
  
    if Error_level_cut_off is not None:
        Error_level_cut_off = float(Error_level_cut_off)

    # Create dataframe
    analyte_dataSummary = pd.concat([x, y, z], axis=1).reset_index(drop=True)
    analyte_dataSummary.index = range(1, len(analyte_dataSummary) + 1)
    analyte_dataSummary.index.name = 'Sample ID'
    analyte_dataSummary = analyte_dataSummary.reset_index()

    # Convert to numeric, coerce errors to NaN, because we perform numeric calculations
    analyte_dataSummary.iloc[:,1] = pd.to_numeric(analyte_dataSummary.iloc[:,1], errors='coerce') #reference method
    analyte_dataSummary.iloc[:,2] = pd.to_numeric(analyte_dataSummary.iloc[:,2], errors='coerce') #test method

    # Calculate differences and percent differences
    # Replace 0 with NaN in the denominator to prevent ZeroDivisionError
    safe_denominator = analyte_dataSummary.iloc[:,1].replace(0, float('nan'))
    
    analyte_dataSummary["Y-X"] = analyte_dataSummary.iloc[:,2] - analyte_dataSummary.iloc[:,1]
    analyte_dataSummary["%(Y-X)/X"] = (analyte_dataSummary["Y-X"] / safe_denominator) * 100
    analyte_dataSummary["Y-X"] = pd.to_numeric(analyte_dataSummary["Y-X"], errors='coerce')
    analyte_dataSummary["%(Y-X)/X"] = pd.to_numeric(analyte_dataSummary["%(Y-X)/X"], errors='coerce')
    analyte_dataSummary[["Y-X", "%(Y-X)/X"]] = analyte_dataSummary[["Y-X", "%(Y-X)/X"]].round(3)

    # Apply Pass/Fail criteria
    def check_pass_fail(row):
        # Check for NaN in x.name or Test Method (Y)
        if pd.isna(row[x.name]) or pd.isna(row[y.name]):
            return "Not assessed"

        if Error_level_cut_off is not None:
            # We removed the float cast from here because we already did it at the top of the main function!
            
            if error1 is None or error2 is None:
                raise ValueError(
                    "Please provide both 'error1' and 'error2' if and when you specify 'Error_level_cut_off' value."
                )
            
            # Check if x is numeric before comparing to the float cut-off
            if isinstance(row[x.name], (int, float)):
                if row[x.name] >= Error_level_cut_off:
                    return "**FAIL**" if abs(row["%(Y-X)/X"]) > error2 else "pass"
                else:
                    return "**FAIL**" if abs(row["Y-X"]) > error1 else "pass"
            else:
                # If it's a string (e.g., "<0.05"), it skips the >= comparison entirely
                return "Not assessed" 

        elif Error_level_cut_off is None:
            if error1 is not None and error2 is None:
                # Check if x is numeric before comparing to the float cut-off
                if isinstance(row[x.name], (int, float)):
                    return "**FAIL**" if abs(row["Y-X"]) > error1 else "pass"
                else:
                    return "Not assessed"  # If it's a string (e.g., "<0.05"), it skips the numeric comparison entirely
            elif error2 is not None and error1 is None:
                if isinstance(row[x.name], (int, float)):
                    return "**FAIL**" if abs(row["%(Y-X)/X"]) > error2 else "pass"
                else:
                    return "Not assessed"
            else:
                return "Not assessed"  # When both error1 and error2 are None or both are provided (not handled here)
                
        return "Not assessed"

    analyte_dataSummary["Pass/Fail"] = analyte_dataSummary.apply(check_pass_fail, axis=1)

    df_selected_columns = analyte_dataSummary.iloc[:,1:3] #select only reference and test method columns
    df_selected_columns["Y-X"] = analyte_dataSummary["Y-X"]
    df_selected_columns["%(Y-X)/X"] = analyte_dataSummary["%(Y-X)/X"]

    # Return sorted dataframe
    return analyte_dataSummary.sort_values(by=[x.name]), df_selected_columns

def Histogram_grouped(x,y):
    """_summary_
        Create a grouped histogram.  This is specifically for looking at how the 
        two methods of measures are distributed
    Args:
        df (data frame): This data frame must contain 'Primary Refence Method Result' and 'Test Method Result' columns.

    Returns:
        _type_: _plotly figure
    """
    # Create the initial DataFrame
    df = pd.DataFrame({'Reference Method': x, 'Test Method': y})

    # Reshape the DataFrame to long format
    df_long = pd.melt(df, 
                      value_vars=['Reference Method', 'Test Method'], 
                      var_name='Method', 
                      value_name='Result')

    # Create the histogram with grouped bars
    fig = px.histogram(df_long, 
                       x='Result',
                       color='Method',
                       histnorm='probability',  # Normalize to show count proportions
                       barmode='group',        # Display bars side by side
                       marginal="box",         # Keep the rug plot for individual data points
                       hover_data=df_long.columns)

    # Update layout for better visualization
    fig.update_layout(
        plot_bgcolor="white",
            paper_bgcolor="white",
            xaxis=dict(showgrid=True, gridcolor="lightgrey"),
            yaxis=dict(showgrid=True, gridcolor="lightgrey"),
        xaxis_title='Result',
        yaxis_title='Count (Normalized)',
        showlegend=True,
        bargap=0.2,  # Adjust gap between bars for clarity
    ) 
    fig.show()
    fig.write_image("Figures/boxplot.png", height = 450, width = 900, scale=2)
    FigWhiteCrop("Figures/boxplot.png")
    
    return fig

def RI_histogram(df, low_lim_RI, high_lim_RI):

    '''
    Must have a data frame with "Result" and "Gender" columns
    Gender must have values as "Male" and/or "Female"
    '''
    # Initialize a dictionary to store medians
    medians = {}

    # Create the histogram
    fig = px.histogram(df, 
                    x='Result',
                    color='Gender',
                    histnorm='',  # Normalize to show count proportions
                    barmode='group',  # Display bars side by side
                    marginal="box",   # Keep the rug plot for individual data points
                    color_discrete_map={'Male': 'lightblue', 'Female': 'pink'},  # Set custom colors
                    )

    # Update layout for better visualization
    fig.update_layout(
        template="plotly_white",
        xaxis_title='Result',
        yaxis_title='Count',
        showlegend=True,
        bargap=0.2,  # Adjust gap between bars for clarity
    )

    # Add lower RI limit
    fig.add_trace(
        go.Scatter(x = [low_lim_RI, low_lim_RI],
                y = [0,len(df)/2 ],
                    mode = "lines",
                    line = dict(color = "black", dash = "dash", width = 2),
                    name = f"Lower RI Limit: {low_lim_RI}" 
                    )
            )
    # Add upper RI limit
    fig.add_trace(
        go.Scatter(x = [high_lim_RI, high_lim_RI],
                y = [0,len(df)/2 ],
                    mode = "lines",
                    line = dict(color = "black", dash = "dash", width = 2),
                    name = f"Lower RI Limit: {high_lim_RI}" 
                    )
            )

    # Compute and add density curves, and store medians
    for gender in df['Gender'].unique():
        # Filter data for the current gender
        data = df[df['Gender'] == gender]['Result'].values
        # Compute and store the median
        median_value = np.median(data)
        medians[gender] = median_value  # Store in dictionary
        print(f"The median for {gender} is {median_value:.2f}")  # Print with 2 decimal places
        medians_df = pd.DataFrame(list(medians.items()), columns=['Gender', 'Median'])
        
        
        countRI = len(data)  # Number of data points
        # Compute KDE using scipy, use the gaussian distribution to smooth the data
        kde = gaussian_kde(data, bw_method=0.5)  # Adjust bw_method for smoothness
        # Generate x values for the KDE curve (spanning the range of Result)
        x_min, x_max = data.min(), data.max()
        kde_x = np.linspace(x_min, x_max, 100)  # 100 points for smooth curve
        kde_y = kde(kde_x)  # Compute KDE y values
        
        # Scale KDE to match histogram height (approximate scaling)
        kde_max = kde_y.max()  # Max value of KDE
        kde_y_scaled = kde_y / (kde_y.max()) * countRI  # Scale KDE to fit histogram
        
        # Add density curve to the plot
        fig.add_trace(
            go.Scatter(
                x=kde_x,
                y=kde_y_scaled,
                mode='lines',
                name=f'{gender} Density',
                line=dict(color='lightblue' if gender == 'Male' else 'pink', width=2)
            )
        )
    
    fig.write_image("Figures/RI.png",  width = 900, height = 500, scale=3)
    FigWhiteCrop("Figures/RI.png")

    fig.show()
    return medians_df


#######################################
# Deming Regression with RSD as variance
#######################################
"""
#Two different flavors or Deming Regression.  These are functions I wrote based on the papers by Linnet et al in Clinical Chemistry
#1. with lambda set calcluation of CV or the RSD of the methods ^2
    1. has 4 input variables

        x = Reference Method
        y = Test Method
        z = Hue; for plotting purposes, which is to define a variable to colour points
        z2 = Style; for plotting purposes, which is to define a variable to style the points
    In this method, above, I made it so that you can enter just x and y, or additional z and z2 arguements.

#2. With lambda =1; this is how several methods calculate the Deming regression, but do not specify
This is listed further down below
"""


# 1. A  Deming Regression with Lambda ratio as RSD for each respective method
def Deming_Method_NonEqual(x, y):
    if len(x) == len(y):
        x_bar = np.mean(x)
        y_bar = np.mean(y)
        p = sum([(xi - x_bar) * (yi - y_bar) for xi, yi in zip(x, y)])
        u = sum((xi - x_bar) ** 2 for xi in x)
        q = sum((yi - y_bar) ** 2 for yi in y)
        # here we take the lambda ratio as (CV_x/CV_y)^2
        lambda_ratio = ((np.std(x) / np.mean(x)) / (np.std(y) / np.mean(y))) ** 2
        # slope
        slope = (
            (lambda_ratio * q - u)
            + (((u - lambda_ratio * q) ** 2) + 4 * lambda_ratio * (p**2)) ** 0.5
        ) / (2 * lambda_ratio * p)
        # intercept
        yintercept = y_bar - slope * x_bar
        # Pearson R
        Pearson_r = np.corrcoef(x, y)[0, 1]
    else:
        print("X and Y Variables are note the same length")
    return slope, yintercept, Pearson_r


# 2. Deming Regression with Equal variance
def Deming_Method(x, y):
    """
    This is the most common type of Deming Regression with Equal Variance
    Here the Lambda ratio is simply set to 1

    """
    if len(x) == len(y):
        x_bar = np.mean(x)
        y_bar = np.mean(y)
        p = sum([(xi - x_bar) * (yi - y_bar) for xi, yi in zip(x, y)])
        u = sum((xi - x_bar) ** 2 for xi in x)
        q = sum((yi - y_bar) ** 2 for yi in y)
        # here we take the lambda ratio as (1) by assuming equal variance
        lambda_ratio = 1
        # slope
        slope = (
            (lambda_ratio * q - u)
            + (((u - lambda_ratio * q) ** 2) + 4 * lambda_ratio * (p**2)) ** 0.5
        ) / (2 * lambda_ratio * p)
        # intercept
        yintercept = y_bar - slope * x_bar
        # Pearson R
        Pearson_r = np.corrcoef(x, y)[0, 1]
    else:
        print("X and Y Variables are note the same length")
    return slope, yintercept, Pearson_r

def reframe_data(x, y, z=None, z2=None):
    '''
    This function takes in x, y, z, and z2 as inputs and returns a cleaned DataFrame  
    with numeric values for x and y, and drops any rows with NaN values.
    
    z and z2 are optional and can be included if provided.
    '''
    # Create a dictionary to hold the data    
    data = {x.name: pd.to_numeric(x, errors="coerce"),
            y.name: pd.to_numeric(y, errors="coerce")}
    if z is not None:
        data[z.name] = z
    if z2 is not None:
        data[z2.name] = z2
        
    return pd.DataFrame(data).dropna()

def Deming_Plot_Equal_Variance_with_Error2_PX(x, y, z=None, z2=None, **error_info):
    """
    Creates a Deming regression plot with error bands using Plotly Express.
    
    Parameters:
    - x, y: Pandas Series or array-like, numeric data for x and y axes.
    - z: Continuous variable for color mapping (optional).
    - z2: Categorical variable for symbol mapping (optional).
    - error_info: Dictionary with keys 'error1' (absolute error), 'error2' (percentage error),
                  and 'Error_level_cut_off' (concentration cut-off).

    Returns:
    - Tuple: (number of points, slope, y-intercept, Pearson R).
    """
    # Extract error parameters
    error1 = error_info.get("error1")
    error2 = error_info.get("error2")
    cutoff = error_info.get("Error_level_cut_off")
    output_path = error_info.get("output_path", "Figures/Deming_regression.png")

    # Reframe and clean data
    df = reframe_data(x, y, z, z2)

    # Extract cleaned data
    x, y = df[x.name], df[y.name]
    z = df[z.name] if z is not None else None
    z2 = df[z2.name] if z2 is not None else None
    xlen = len(x)

    # Perform Deming regression
    slope, yintercept, pearson_r = Deming_Method(x, y)

    # Create error band x-values
    x_range = np.linspace(x.min(), x.max(), 100)
    
    # Calculate error bands
    if error1 is None and error2 is None:
        yabove = ybelow = x_range
    elif error1 is not None and error2 is None:
        yabove = x_range + error1
        ybelow = x_range - error1
    elif error2 is not None and error1 is None:
        yabove = x_range * (1 + error2 / 100)
        ybelow = x_range * (1 - error2 / 100)
    else:
        yabove = np.where(x_range <= cutoff, x_range + error1, x_range * (1 + error2 / 100))
        ybelow = np.where(x_range <= cutoff, x_range - error1, x_range * (1 - error2 / 100))

    scatter_args = {"x": x, "y": y}  # base/mandatory arguments
    if z is not None and z2 is not None:
        scatter_args.update({"color": z, "symbol": z2, "labels": {"color": z.name, "symbol": z2.name}})
    elif z is not None:
        scatter_args.update({"color": z, "labels": {"color": z.name}})
    elif z2 is not None:
        scatter_args.update({"symbol": z2, "labels": {"symbol": z2.name}})
    scatter_fig = px.scatter(**scatter_args, color_continuous_scale="Bluered_r")

    # Initialize figure
    fig = go.Figure()

    # Add error band (first layer)
    fig.add_trace(
        go.Scatter(
            x=np.concatenate([x_range, x_range[::-1]]),
            y=np.concatenate([yabove, ybelow[::-1]]),
            fill="toself",
            fillcolor="rgba(211,211,211,0.60)",  # dark gray with 15% opacity
            line=dict(color="rgba(255,255,255,0)"),
            name="Allowable Error",
            showlegend=True,
        )
    )

    # Add scatter traces, ensuring legend and color scale
    for trace in scatter_fig.data:

        if not trace.name: #If the trace doesn't have a name (e.g., when z and z2 are None)
            trace.name = "Samples" # give it a default label
            
        trace.update(showlegend=True)  # Ensure legend is shown for z2 and z, i.e. use labels above.
        fig.add_trace(trace)

    # Apply color scale to final figure if z is provided
    if z is not None:   
        fig.update_layout(coloraxis=dict(colorscale="Bluered_r"))
    
    # Add line of identity
    expansion = 0.025 * x.min()
    x_ext = [x.min() - expansion, x.max() + expansion]
    fig.add_trace(
        go.Scatter(
            x=x_ext,
            y=x_ext,
            mode="lines",
            line=dict(color="black", dash="dot", width=0.8),
            name="Line of Identity",
        )
    )

    # Add Deming regression line
    fig.add_trace(
        go.Scatter(
            x=x_ext,
            y=[slope * x_ext[0] + yintercept, slope * x_ext[1] + yintercept],
            mode="lines",
            line=dict(color="Purple", dash="solid", width=1.0),
            name="Line of Fit",
        )
    )

    # Add regression annotation
    fig.add_annotation(
        xref="paper",
        yref="paper",
        x=0.1,
        y=0.9,
        xanchor="left",
        yanchor="top",
        text=f"Regression by Deming<br>y = {slope:.2f}x{'+' if yintercept > 0 else ''}{yintercept:.2f}<br>Pearson R: {pearson_r:.3f}<br>n = {xlen}",
        showarrow=False,
        font=dict(size=12, color="black"),
    )

    # Update layout
    fig.update_layout(
        #plot_bgcolor="white",
        #paper_bgcolor="white",
        xaxis=dict(showgrid=True, gridcolor="lightgrey", title=x.name),
        yaxis=dict(showgrid=True, gridcolor="lightgrey", title=y.name),
        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="center", x=0.5),
        template="plotly_white",
        width=600,
        height=600,
        )

    # Save figure (optional, configurable)
    fig.write_image(output_path, height=600, width=600, scale=2)
    FigWhiteCrop("Figures/Deming_regression.png")
    # Show figure
    fig.show()
    #print(f"Deming regression plot saved to {output_path}")
    return xlen, slope, yintercept, pearson_r

#####################################################
# Difference plot with Median And error regions ####
########################################################

def Difference_plot_median_with_error2(x, y, z=None, z2=None, **error_info):
    """
    error1 = error in absolute terms as indicated in IQMH
    error2 = error in percentage terms as indicated in IQMH; not considered in this function and plot
    concentration_cut = the cut-off at which error1 and error2 are applicable

    if concentration_cut is not provided, but only one of the error terms is provided, 
    then error (either error1 or erro2) is calculated across all range of x. 
    """
    # Extract parameters
    error1 = error_info.get("error1", None)
    Error_level_cut_off = error_info.get("Error_level_cut_off", None)
    output_path = error_info.get("output_path", "Figures/Difference_Plot_median.png")

    # Input validation
    if not hasattr(x, 'name') or not hasattr(y, 'name'):
        raise ValueError("Inputs x and y must have a 'name' attribute (e.g., pandas Series).")
    if z is not None and not hasattr(z, 'name'):
        raise ValueError("Input z must have a 'name' attribute if provided.")
    if z2 is not None and not hasattr(z2, 'name'):
        raise ValueError("Input z2 must have a 'name' attribute if provided.")

    # Reframe and clean data
    df = reframe_data(x, y, z, z2)
    x, y = df[x.name], df[y.name]
    z = df[z.name] if z is not None else None
    z2 = df[z2.name] if z2 is not None else None
    ydiff = y - x

    # Build scatter plot arguments dynamically
    scatter_args = {
        "x": x,
        "y": ydiff,
        "labels": {
            "x": x.name,
            "y": f"{y.name} - {x.name}"
        }
    }
    if z is not None:
        scatter_args["color"] = z
        scatter_args["color_continuous_scale"] = "Bluered_r"
        scatter_args["labels"]["color"] = z.name
    if z2 is not None:
        scatter_args["symbol"] = z2
        scatter_args["labels"]["symbol"] = z2.name

    fig = px.scatter(**scatter_args)

    # Calculate expansion factor
    expansion_factor = 0.025 * (x.max() - x.min())

    # x-values to span with some expansion
    x_range=[x.min() - expansion_factor, x.max() + expansion_factor]
    
    # median bias line
    median_y = ydiff.median()
    fig.add_trace(
        go.Scatter(
            x = x_range,
            y=[median_y, median_y],
            mode="lines",
            line=dict(color="red", dash="solid", width=0.8),
            name=f"Median Bias: {round(median_y, 1)}",
        )
    )

    #plot IQR lines
    q1 = ydiff.quantile(0.25)
    q3 = ydiff.quantile(0.75)
    for q_val, label in zip([q1, q3], ["25th Pct", "75th Pct"]):
        fig.add_trace(go.Scatter(
            x=x_range, y=[q_val, q_val],
            mode="lines", line=dict(color="red", dash="dot", width=0.6),
            name=f"{label}: {round(q_val, 2)}"
        ))

    # Add agreement line or the zero line
    fig.add_trace(
        go.Scatter(
            x=x_range,
            y=[0, 0],
            mode="lines",
            line=dict(color="black", dash="dot", width=0.8),
            name="Agreement Line",
        )
    )

    # Add error band if error1 is provided
    if error1:
        x1 = Error_level_cut_off if Error_level_cut_off not in (0, None) else x.max() + expansion_factor
        fig.add_shape(
            type="rect",
            x0=x.min() - expansion_factor,
            x1=x1,
            y0=0 + error1,
            y1=0 - error1,
            fillcolor="lightgrey", 
            opacity=0.6,
            layer="below",
            line_width=0,
            name ="Allowable Error",
        )

    # Update layout
    fig.update_layout(
        plot_bgcolor="white",
        paper_bgcolor="white",
        xaxis=dict(showgrid=True, gridcolor="lightgrey", title=x.name),
        yaxis=dict(showgrid=True, gridcolor="lightgrey", title=f"{y.name} - {x.name}"),
        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="center", x=0.5),
    )

    # Save and crop figure
    fig.show()
    fig.write_image(output_path, scale=2)
    FigWhiteCrop(output_path)

    return median_y

#####################################################
# % Difference plot with Median And Error Regions ####
########################################################

def Difference_plot_Percent_median_with_error2(x, y, z=None, z2=None, **error_info):
    """
    error1 = error in absolute terms as indicated in IQMH - Not considered in this function and plot
    error2 = error in percentage terms as indicated in IQMH
    concentration_cut = the cut-off at which error1 and error2 are applicable

    If concentration_cut is not provided, but only one of the error terms is provided, 
    then error (either error1 or erro2) is calculated across all range of x.
    """

    # Extract parameters
    error2 = error_info.get("error2", None)
    Error_level_cut_off = error_info.get("Error_level_cut_off", None)
    output_path = error_info.get("output_path", "Figures/Difference_Plot_median_per.png")

    # Input validation
    if not hasattr(x, 'name') or not hasattr(y, 'name'):
        raise ValueError("Inputs x and y must have a 'name' attribute (e.g., pandas Series).")
    if z is not None and not hasattr(z, 'name'):
        raise ValueError("Input z must have a 'name' attribute if provided.")
    if z2 is not None and not hasattr(z2, 'name'):
        raise ValueError("Input z2 must have a 'name' attribute if provided.")

    # Convert to numeric and create DataFrame
    df = reframe_data(x, y, z, z2)
    # Extract cleaned data
    x, y = df[x.name], df[y.name]
    z = df[z.name] if z is not None else None
    z2 = df[z2.name] if z2 is not None else None

    # Calculate percent difference
    # Replace 0 with NaN in x to prevent zero-division without aborting the script
    x = x.replace(0, float('nan'))
    ydiff = 100 * (y - x) / x

    # Build scatter plot arguments dynamically
    scatter_args = {
        "x": x,
        "y": ydiff,
        "labels": {
            "x": x.name,
            "y": "% Difference"
        }
    }
    if z is not None:
        scatter_args["color"] = z
        scatter_args["color_continuous_scale"] = "Bluered_r"
        scatter_args["labels"]["color"] = z.name
    if z2 is not None:
        scatter_args["symbol"] = z2
        scatter_args["labels"]["symbol"] = z2.name
    fig = px.scatter(**scatter_args)

    # Calculate expansion factor
    expansion_factor = 0.025 * (x.max() - x.min())
    x_range=[x.min() - expansion_factor, x.max() + expansion_factor]
    
    # Add median bias line
    median_y = ydiff.median()
    fig.add_trace(
        go.Scatter(
            x=x_range,
            y=[median_y, median_y],
            mode="lines",
            line=dict(color="red", dash="solid", width=0.8),
            name=f"Median Bias: {round(median_y, 1)}%",
        )
    )

    #plot IQR lines
    q1 = ydiff.quantile(0.25)
    q3 = ydiff.quantile(0.75)
    for q_val, label in zip([q1, q3], ["25th Pct", "75th Pct"]):
        fig.add_trace(go.Scatter(
            x=x_range, y=[q_val, q_val],
            mode="lines", line=dict(color="red", dash="dot", width=0.6),
            name=f"{label}: {round(q_val, 2)}%"
        ))

    # Add agreement line
    fig.add_trace(
        go.Scatter(
            x=[x.min() - expansion_factor, x.max() + expansion_factor],
            y=[0, 0],
            mode="lines",
            line=dict(color="black", dash="dot", width=0.8),
            name="Agreement Line",
        )
    )

    # Add error band if error2 is provided
    if error2:
        if Error_level_cut_off not in (0, None):
            x0 = Error_level_cut_off
            x1 = x.max() + expansion_factor
        else:
            x0 = x.min() - expansion_factor
            x1 = x.max() + expansion_factor
        fig.add_shape(
            type="rect",
            x0=x0,
            x1=x1,
            y0=0 + error2,
            y1=0 - error2,
            fillcolor="lightgrey",
            opacity=0.6,
            layer="below",
            line_width=0,
            name = "Allowable Error"
        )

    # Update layout
    fig.update_layout(
        plot_bgcolor="white",
        paper_bgcolor="white",
        xaxis=dict(showgrid=True, gridcolor="lightgrey", title=x.name),
        yaxis=dict(showgrid=True, gridcolor="lightgrey", title="% Difference"),
        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="center", x=0.5),
    )

    # Save and crop figure
    fig.show()
    fig.write_image(output_path, scale=2)
    FigWhiteCrop(output_path)

    return median_y
##
# Function to write results to CSV
def write_to_csv(analyte, n, slope, y_intercept, pearson_r, median_bias, percent_bias, unit, filename="MethodSummary.csv"):
    '''
    Write the results to a CSV file.
    
    Args:
        analyte: The analyte name or identifier.
        n: Sample size.
            Obtainble from function Deming_Plot_Equal_Variance_with_Error2_PX
        slope: Slope of the regression line.
            Obtainble from function Deming_Plot_Equal_Variance_with_Error2_PX
        y_intercept: Y-intercept of the regression line.
            Obtainble from function Deming_Plot_Equal_Variance_with_Error2_PX
        pearson_r: Pearson correlation coefficient from function.
            Obtainble from function Deming_Plot_Equal_Variance_with_Error2_PX
        median_bias: Median bias value.
            Obtainble from function Difference_plot_median_with_error2
        percent_bias: Percent bias value.
            Obtainble from function Difference_plot_Percent_median_with_error2
        unit: Unit of measurement.
        filename: Name of the CSV file (default: 'results.csv').
    '''
    
    # Define the header for the CSV file
    header = ["Analyte", "N", 'Slope', 'y-Intercept', "Pearson's r", 'Median Bias', '% Bias', 'Unit']
    
    # Check if file exists to determine if we need to write the header
    file_exists = os.path.isfile(filename)
    
    # Open the file in append mode
    with open(filename, mode='a', newline='') as file:
        writer = csv.writer(file)
        
        # Write header if file doesn't exist
        if not file_exists:
            writer.writerow(header)
            
        # Write the data row
        writer.writerow([analyte, n, slope, y_intercept, pearson_r, median_bias, percent_bias, unit])

# function to create a word table from our data frame
# You have this function below else where.  Consider make this function as part of a class
def create_word_table(document, MyTable):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    table = document.add_table(MyTable.shape[0] + 1, MyTable.shape[1])
    table.style = document.styles["Light Shading Accent 1"] # previously set to "Light Grid Accent 1"
    table.autofit = True
    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # Set column widths (optional) 
    for col in table.columns:
        for cell in col.cells:
            cell.width = Cm(2.5)  # Set each column width to 3 cm
    

    # Add the header rows
    for j in range(MyTable.shape[-1]):
        cell = table.cell(0, j)
        cell.text = MyTable.columns[j]
        # Set font size for header cell
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # center text in each cell
            for run in paragraph.runs:
                run.font.size = Pt(10)
    
    # Add the rest of the data frame
    for i in range(MyTable.shape[0]):
        for j in range(MyTable.shape[-1]):
            cell = table.cell(i + 1, j)
            cell.text = str(MyTable.values[i, j])
            # Set font size for data cell
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # center text in each cell
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    document.save("MethodComparison.docx")


"""
Precision Studies - These are based on using the ANOVA approach.

I checked these against the traditional method that did not use the ANOVA approach.
Earlier CLSI guidelines did not use ANOVA approach. But the most recent one does
I showed that both methods are equivalent when we evaluated the sTfR method using the Randox Reagnet.
"""
##
# this function is for getting a summary table of the precision or methods
# Essentially describes the column data using discriptive stats
# Also computes the CV
# get only the columns that are not-strings, i.e. methods onlys
# One can check the results of this with 'df.describe()'

def data_summary(dataAll):
    """
    Returns a professional summary table with automatic, correct significant figures
    using the gold-standard 'uncertainties' package.
    
    Mean ± SEM is displayed perfectly formatted with proper scientific rounding.
    """
    # Keep only numeric columns
    df = dataAll.select_dtypes(include=[np.number])
    
    # We'll build a list of dicts → clean, readable, and flexible
    rows = []
    
    for col in df.columns:
        data = df[col].dropna()
        n = len(data)
        
        if n == 0:
            continue
        if n == 1:
            # Only one value → SEM undefined
            mean = data.iloc[0]
            sem = np.nan
            sd = np.nan
            cv = np.nan
        else:
            mean = data.mean()
            sd = data.std(ddof=1)
            sem = sd / np.sqrt(n)
            cv = round((sd / mean * 100),2) if mean != 0 else np.nan
        
        # This is the magic line — ufloat handles ALL sig fig rules automatically
        mean_with_unc = ufloat(mean, sem) if n > 1 else mean
        
        rows.append({
            "Variable": col,
            "N": int(n),
            "Mean ± SEM": mean_with_unc,          
            "Mean": mean,
            "SEM": sem,
            "SD": sd,
            "Min": data.min(),
            "Median": data.median(),
            "Max": data.max(),
            "%CV": cv
        })
    
    summary = pd.DataFrame(rows)
    
    # Reorder columns nicely
    col_order = ["Variable", "N", "Mean ± SEM", "Mean", "SEM", "SD", "Min", "Median", "Max", "%CV"]
    summary = summary[col_order]
    
    # Optional: round numeric columns that aren't already perfect
    numeric_cols = ["Mean", "SEM", "SD", "%CV", "Min", "Median", "Max"]
    summary[numeric_cols] = summary[numeric_cols].round(6)  # safe, won't hurt ufloat display
    
    return summary
    

########################################################
############### LEVEY-JENNINGS (L-J Chart)   ###########

def LJ_plot2(df, x, y, Replicate=None, QCRun=None):
        
    fig = px.scatter(df, x=x, y=y, color=QCRun, symbol=Replicate, labels=('Replicate', 'Run'))
    mean_y = y.mean()
    sd_y = y.std()

    # shade the region between +1SD and 2SD
    fig.add_shape(
        type="rect",
        x0=0,
        x1=x.max(),
        y0=mean_y + sd_y,
        y1=mean_y + 2 * sd_y,
        fillcolor="lightgrey",
        opacity=0.5,
        layer="below",
        line_width=0,
    )
    # shade the region between -1SD and -2SD
    fig.add_shape(
        type="rect",
        x0=0,
        x1=x.max(),
        y0=mean_y - sd_y,
        y1=mean_y - 2 * sd_y,
        fillcolor="lightgrey",
        opacity=0.5,
        layer="below",
        line_width=0,
    )
    # shade the region between -2SD and -3SD
    fig.add_shape(
        type="rect",
        x0=0,
        x1=x.max(),
        y0=mean_y - 2 * sd_y,
        y1=mean_y - 3 * sd_y,
        fillcolor="lightgrey",
        opacity=0.9,
        layer="below",
        line_width=0,
    )
    # shade the region between +2SD and 3SD
    fig.add_shape(
        type="rect",
        x0=0,
        x1=x.max(),
        y0=mean_y + 2 * sd_y,
        y1=mean_y + 3 * sd_y,
        fillcolor="lightgrey",
        opacity=0.9,
        layer="below",
        line_width=0,
    )
    # Add a horizontal line at the mean y-value with a specified line width
    fig.add_trace(
        go.Scatter(
            x=[0, x.max()],
            y=[mean_y, mean_y],
            mode="lines",
            line=go.scatter.Line(color="black", dash="solid", width=0.8),
            name="Mean",
        )
    )
    # Add a horizontal line at the +1SD with a specified line width
    fig.add_trace(
        go.Scatter(
            x=[0, x.max()],
            y=[mean_y + sd_y, mean_y + sd_y],
            mode="lines",
            line=go.scatter.Line(color="black", dash="dash", width=0.8),
            name="+1 SD",
        )
    )
    # Add a horizontal line at the -1SD with a specified line width
    fig.add_trace(
        go.Scatter(
            x=[0, x.max()],
            y=[mean_y - sd_y, mean_y - sd_y],
            mode="lines",
            line=go.scatter.Line(color="black", dash="dash", width=0.8),
            name="- 1 SD",
        )
    )
    # Add a horizontal line at the +2SD with a specified line width
    fig.add_trace(
        go.Scatter(
            x=[0, x.max()],
            y=[mean_y + 2 * sd_y, mean_y + 2 * sd_y],
            mode="lines",
            line=go.scatter.Line(color="red", dash="dash", width=0.8),
            name="+2 SD",
        )
    )
    # Add a horizontal line at the -2SD with a specified line width
    fig.add_trace(
        go.Scatter(
            x=[0, x.max()],
            y=[mean_y - 2 * sd_y, mean_y - 2 * sd_y],
            mode="lines",
            line=go.scatter.Line(color="red", dash="dash", width=0.8),
            name="- 2 SD",
        )
    )

    # Add a horizontal line at the -3SD with a specified line width
    fig.add_trace(
        go.Scatter(
            x=[0, x.max()],
            y=[mean_y - 3 * sd_y, mean_y - 3 * sd_y],
            mode="lines",
            line=go.scatter.Line(color="red", dash="dash", width=0.8),
            name="- 3 SD",
        )
    )
    # Add a horizontal line at the +3SD with a specified line width
    fig.add_trace(
        go.Scatter(
            x=[0, x.max()],
            y=[mean_y + 3 * sd_y, mean_y + 3 * sd_y],
            mode="lines",
            line=go.scatter.Line(color="red", dash="dash", width=0.8),
            name="+3 SD",
        )
    )

    # update the layout to white background
    fig.update_layout(
        #plot_bgcolor="white",
        #paper_bgcolor="white",
        xaxis=dict(showgrid=True, gridcolor="lightgrey"),
        yaxis=dict(showgrid=True, gridcolor="lightgrey"),
        template = "plotly_white",
        width = 700,
        height = 500
    )
    if Replicate is None:
        fig.update_layout(legend=dict(title=dict(text="Run")))
    if Replicate is not None:
        fig.update_layout(legend=dict(title=dict(text="Run, Replicate")))
        
    fig.show()
    fig.write_image("Figures/LJ-Plot.png", width = 700, height = 500, scale=2)
    FigWhiteCrop("Figures/LJ-Plot.png")



def tukey_outliers_remove(col_name, df):
    '''Function to apply Tukey's method for outlier detection
    Args:
        col_name (str): Name of the column to check for outliers
        df (DataFrame): Pandas DataFrame containing the data
    Returns:
        dataframe with outliers removed
    '''
    column = df[col_name]
    
    #Calculate Q1 (25th percentile) and Q3 (75th percentile)
    Q1 = df[col_name].quantile(0.25)
    Q3 = df[col_name].quantile(0.75)
    IQR = Q3 - Q1
    
    #Determine outliers using 1.5*IQR rule, i.e. lower and upper fences/bounds
    lower_bound = Q1 - 1.5 * IQR
    upper_bound = Q3 + 1.5 * IQR
    print(f"Lower Bound: {lower_bound}, Upper Bound: {upper_bound}")
    print(f"Number of outliers removed: {len(column[(column < lower_bound) | (column > upper_bound)])}")
    print("Below are the outliers (first 10 rows):")
    #outliers = df[[col_name]][(column < lower_bound) | (column > upper_bound)]
    outliers = df[(column < lower_bound) | (column > upper_bound)]
    outliers = outliers.sort_values(by=col_name)
    print(outliers.head(10))
    print('')
    print(f"Number of values retained: {len(column[(column >= lower_bound) & (column <= upper_bound)])}")
    print("Below are the values retained (first 10 rows):")
    #df_retained = df[[col_name]][(column >= lower_bound) & (column <= upper_bound)]
    df_retained = df[(column >= lower_bound) & (column <= upper_bound)]
    print(df_retained.sort_values(by=col_name).head(10))
    print('')
    print('summary of data retained:')
    print(df_retained.describe())
    return df_retained


#######################################################
####### Linearity Plots ##############################
#######################################################
########################################################

def linearity_plots(Measured_mean,Expected_mean, Measured_err, n, **error_info):
    """
    Perform linearity analysis and generate regression and residual plots.

    Parameters:

    Measured_mean (pd.Series): Series of measured mean values.
    Expected_mean (pd.Series): Series of expected mean values.
    Error_level_cut_off (float): Cut-off value for allowable error.
    error1 (float): Absolute error value for levels below the cut-off.
    error2 (float): Percentage error value for levels above the cut-off.

    Returns:
    fig: Plotly figure object containing the linearity and residuals plots.
    """
    
    error1 = error_info.get("error1", None)
    error2 = error_info.get("error2", None)
    Error_level_cut_off = error_info.get("Error_level_cut_off", None)
    
    x = Expected_mean
    y = Measured_mean
    slope, intercept, r_value, p_value, std_err = linregress(x, y)
    r_squared = r_value**2
    equation = f"y = {slope:.3f}x + {intercept:.3f}"
    xmin = x.min()
    xmax = x.max()
    ymin = y.min()
    ymax = y.max()
    xfit = [xmin, xmax]
    yfit = [slope * xmin + intercept, slope * xmax + intercept]

    # Create the error bands
    if (x.max() - x.min()) == 0:
        x_er = np.array([x.min()])
    else:
        x_er = np.arange(min(x), max(x), (max(x) - min(x)) / 100)
    
    if error1 is None and error2 is None:
        # Scenario for when no error limit is provided
        yabove = x_er
        ybelow = x_er
    elif error1 is not None and error2 is None:
        yabove = x_er + error1
        ybelow = x_er - error1
    elif error2 is not None and error1 is None:
        yabove = (1 + error2 / 100) * x_er
        ybelow = (1 - error2 / 100) * x_er
    else:
        yabove = np.where(
            x_er <= Error_level_cut_off, x_er + error1, (1 + error2 / 100) * x_er
        )
        ybelow = np.where(
            x_er <= Error_level_cut_off, x_er - error1, (1 - error2 / 100) * x_er
        )
        
    #Calculate residuals
    residuals = y - (slope * x + intercept)

    #Create a subplot figure with 1 row and 2 columns
    fig = make_subplots(rows=1, cols=2,
                        subplot_titles=("Linearity Plot", "Residuals Plot"),
                        vertical_spacing=0.1)

    # First plot: Linearity
    fig.add_trace(go.Scatter(x=x,
                            y=y,
                            mode='markers',
                            error_y=dict(type='data', array=Measured_err),
                            name='Data Points'),
                            row=1, col=1)

    #Add fitted regression line
    fig.add_trace(go.Scatter(x=xfit,
                            y=yfit,
                            mode='lines',
                            line=dict(color='darkblue', width=0.5),
                            name='Fitted Line'),
                row=1, col=1)

    # Add line of identity
    fig.add_trace(go.Scatter(x = [0, xmax],
                            y = [0, xmax],
                            mode = 'lines',
                            name = 'Line of Identity',
                            line = dict(color = 'red', dash = "dot", width = 2),
                            showlegend=True
                            ),
                            row=1, col=1)

    # Add annotation for equation and R²
    fig.add_annotation(text=f"{equation} <br> R² = {r_squared:.3f}",
                    xref="paper", yref="paper",
                    x=(xmax-xmin)/2, y=ymax-(0.1*(ymax-ymin)),
                    showarrow=False,
                    font=dict(size=12, color = 'darkblue'),
                    align='center',
                    row=1, col=1)
    # Add error band
    
    fig.add_trace(go.Scatter(x=np.concatenate([x_er, x_er[::-1]]),
                            y=np.concatenate([yabove, ybelow[::-1]]),
                            fill="toself",
                            fillcolor="rgba(0, 139, 139, .25)",  # Dark cyan with 25% opacity
                            line=dict(color="rgba(255,255,255,0)"),
                            name="Allowable Error",
                            #legendgroup="error_band",
                            showlegend=True,
                            ), 
                            row = 1, col=1)

    # Second plot: Residuals
    fig.add_trace(go.Scatter(x=x,
                            y=residuals,
                            mode='markers',
                            name='Residuals',
                            ),
                            row=1, col=2)

    # Add horizontal line at y=0
    fig.add_trace(go.Scatter(x=[xmin, xmax],
                            y=[0, 0],
                            mode='lines',
                            line=dict(color='red', dash = "dot",  width=2),
                            name='Zero Line',
                            showlegend=False
                            ),
                            row=1, col=2)

    fig.update_xaxes(title_text="Expected Mean", row=1, col=1)
    fig.update_yaxes(title_text="Measured Mean", row=1, col=1)
    fig.update_xaxes(title_text="Expected Mean", row=1, col=2)
    fig.update_yaxes(title_text="Residuals", row=1, col=2)
    fig.update_layout(
        #title_text="Linearity Analysis",
        template="plotly_white",
        legend=dict(orientation="h", yanchor="bottom", y=1.1, xanchor="center", x=0.5),
        width=900,
        height=500,
        showlegend=True
    )

    fig.write_image("Figures/Linearity.png",  width = 900, height = 500, scale=3)
    FigWhiteCrop2("Figures/Linearity.png")

    fig.show()

"""
Function to graph the a mountain plot of the Percent difference (pd)
between the methods. 
"""
def mountain_plot_pd2(x, y, z=None, frac=0.3):
    """
    Generate a Plotly mountain plot for percent differences.
    Fit the data with a Lowess regression.

    Args:
        x (array-like): Reference method data.
        y (array-like): Test method data.
        z (array-like, optional): Category labels for coloring. If None, uses x.
        frac (float): The fraction of the data used to fitt local regression with Lowess.

    Returns:
        tuple: (plotly.graph_objects.Figure, pd.DataFrame)
            The Plotly figure object and the DataFrame containing the plot data.

    """

    if z is None:
        z = x

    # Create a new DataFrame
    df_Mt = pd.DataFrame(
        {x.name: x, "Test Method (Y)": y, "Color_var": z}
    )

    # Calculate differences and percent differences
    df_Mt["Y-X"] = df_Mt["Test Method (Y)"] - df_Mt[x.name]
    df_Mt["Percent Difference"] = (df_Mt["Y-X"] / df_Mt[x.name]) * 100

    ## Generate DataFrame for % Difference plot
    df_Mt_p_data = df_Mt.sort_values(by="Percent Difference")
    df_Mt_p_data["Rank"] = df_Mt_p_data["Percent Difference"].rank(ascending=True)
    df_Mt_p_data["Percentile"] = df_Mt_p_data["Rank"] * 100 / (len(df_Mt) + 1)

    # Split into above and below 50th percentile
    above50 = df_Mt_p_data[df_Mt_p_data["Percentile"] > 50].copy()
    above50["Percentile"] = 100 - above50["Percentile"]
    below50 = df_Mt_p_data[df_Mt_p_data["Percentile"] <= 50]
    df_pdif_MtPlt = pd.concat([below50, above50])

    # Calculate median bias
    medianBias = round(np.median(df_pdif_MtPlt["Percent Difference"]), 1)

    # Fit the mountain plot data with GAM
    x_fit = df_pdif_MtPlt["Percent Difference"]
    y_fit = df_pdif_MtPlt["Percentile"]
    # frac is the smoothing parameter, frac = 0.3 uses 30% of the data for each LOESS fit
    loess_fit = lowess(y_fit, x_fit, frac=frac)

    # Create the Plotly figure
    fig = px.scatter(
        df_pdif_MtPlt,
        x="Percent Difference",
        y="Percentile",
        color="Color_var",
        color_continuous_scale="Bluered_r",

        labels={
            "Percent Difference": "% Difference (%)",
            "Percentile": "Percentile (%)",
        },
        title="Mountain Plot",
    )

    # Add median bias line
    fig.add_shape(
        type="line",
        x0=medianBias,
        y0=0,
        x1=medianBias,
        y1=50,
        line=dict(color="red", dash="solid", width=1),
        name="Median Bias",
    )

    # Add reference line
    fig.add_shape(
        type="line",
        x0=0,
        y0=0,
        x1=0,
        y1=50,
        line=dict(color="black", dash="dash", width=0.75),
        name="Reference",
    )

    # Annotate median bias
    fig.add_annotation(
        x=medianBias,
        y=52,
        text=f"Median Bias: {medianBias}%",
        showarrow=False,
        font=dict(color="red"),
    )

    # Add LOESS prediction line
    fig.add_trace(
        go.Scatter(
            x=loess_fit[:, 0],
            y=loess_fit[:, 1],
            mode="lines",
            line=dict(color="grey", dash="solid", width=1),
            name="LOESS Prediction",
        )
    )

    fig.update_layout(
        legend=dict(
            title="Color_var",
            orientation="h",
            yanchor="bottom",
            y=1.02,
            xanchor="right",
            x=1,
        )
    )

    return fig, df_pdif_MtPlt

"""
Example usage of the function:
fig, df_pdif_MtPlt = mountain_plot_pd2(x = df_MC['TBIL'],
                  y = df_MC['TBIL2'],
                  z = df_MC['TBIL'],
                  frac = 0.3,)

fig.show()
df_pdif_MtPlt.head(10)

"""

def mountain_plot_ad2(x, y, z=None, frac=0.3):
    """
    Generate a Plotly mountain plot for absolute differences (ad)
    Fit the data with a Lowess regression.

    Args:
        x (array-like): Reference method data.
        y (array-like): Test method data.
        z (array-like, optional): Category labels for coloring. If None, uses x.

    Returns:
        tuple: (plotly.graph_objects.Figure, pd.DataFrame)
            The Plotly figure object and the DataFrame containing the plot data.
    """

    if z is None:
        z = x

    # Create a new DataFrame
    df_Mt = pd.DataFrame(
        {x.name: x, "Test Method (Y)": y, "Color_var": z}
    )

    # Calculate differences and percent differences
    df_Mt["Y-X"] = df_Mt["Test Method (Y)"] - df_Mt[x.name]

    ## Generate DataFrame for Difference plot
    df_Mt_p_data = df_Mt.sort_values(by="Y-X")
    df_Mt_p_data["Rank"] = df_Mt_p_data["Y-X"].rank(ascending=True)
    df_Mt_p_data["Percentile"] = df_Mt_p_data["Rank"] * 100 / (len(df_Mt) + 1)

    # Split into above and below 50th percentile
    above50 = df_Mt_p_data[df_Mt_p_data["Percentile"] > 50].copy()
    above50["Percentile"] = 100 - above50["Percentile"]
    below50 = df_Mt_p_data[df_Mt_p_data["Percentile"] <= 50]
    df_pdif_MtPlt = pd.concat([below50, above50])

    # Calculate median bias
    medianBias = round(np.median(df_pdif_MtPlt["Y-X"]), 1)

    # Fit the mountain plot data with LOESS
    x_fit = df_pdif_MtPlt["Y-X"]
    y_fit = df_pdif_MtPlt["Percentile"]
    # frac is the smoothing parameter, frac = 0.3 uses 30% of the data for each LOESS fit
    loess_fit = lowess(y_fit, x_fit, frac=frac)
    # Create the Plotly figure
    fig = px.scatter(
        df_pdif_MtPlt,
        x="Y-X",
        y="Percentile",
        color="Color_var",
        color_continuous_scale="Bluered_r",
        labels={"Y-X": "Method Difference", "Percentile": "Percentile (%)"},
        title="Mountain Plot",
    )

    # Add median bias line
    fig.add_shape(
        type="line",
        x0=medianBias,
        y0=0,
        x1=medianBias,
        y1=50,
        line=dict(color="red", dash="solid", width=1),
        name="Median Bias",
    )

    # Add reference line
    fig.add_shape(
        type="line",
        x0=0,
        y0=0,
        x1=0,
        y1=50,
        line=dict(color="black", dash="dash", width=0.75),
        name="Reference",
    )

    # Annotate median bias
    fig.add_annotation(
        x=medianBias,
        y=52,
        text=f"Median Bias: {medianBias}",
        showarrow=False,
        font=dict(color="red"),
    )

    # Add LOESS prediction line
    fig.add_trace(
        go.Scatter(
            x=loess_fit[:, 0],
            y=loess_fit[:, 1],
            mode="lines",
            line=dict(color="grey", dash="solid", width=1),
            name="LOESS Prediction",
        )
    )

    fig.update_layout(
        legend=dict(
            title="Color_var",
            orientation="h",
            yanchor="bottom",
            y=1.02,
            xanchor="right",
            x=1,
        )
    )

    return fig, df_pdif_MtPlt


"""
fig, df_pdif_MtPlt = CC.mountain_plot_pd2(x = df_MC['TBIL'],
                    y = df_MC['TBIL2'],
                    z = df_MC['TBIL'],
                    frac=0.3,)  


fig.show()
print(df_pdif_MtPlt.head())

"""
def mountain_plot_pd2_cubic_spline(x, y, z=None, n_splines=12, spline_order=3):
    """
    Generate a Plotly mountain plot for percent differences using Penalized Cubic Regression Splines.

    Args:
        x (array-like): Reference method data.
        y (array-like): Test method data.
        z (array-like, optional): Category labels for coloring. If None, uses x.

    Returns:
        tuple: (plotly.graph_objects.Figure, pd.DataFrame)
            The Plotly figure object and the DataFrame containing the plot data.
    """

    if z is None:
        z = x

    # Create a new DataFrame
    df_Mt = pd.DataFrame(
        {x.name: x, "Test Method (Y)": y, "Color_var": z}
    )

    # Calculate differences and percent differences
    df_Mt["Y-X"] = df_Mt["Test Method (Y)"] - df_Mt[x.name]
    df_Mt["Percent Difference"] = (df_Mt["Y-X"] / df_Mt[x.name]) * 100

    ## Generate DataFrame for % Difference plot
    df_Mt_p_data = df_Mt.sort_values(by="Percent Difference")
    df_Mt_p_data["Rank"] = df_Mt_p_data["Percent Difference"].rank(ascending=True)
    df_Mt_p_data["Percentile"] = df_Mt_p_data["Rank"] * 100 / (len(df_Mt) + 1)

    # Split into above and below 50th percentile
    above50 = df_Mt_p_data[df_Mt_p_data["Percentile"] > 50].copy()
    above50["Percentile"] = 100 - above50["Percentile"]
    below50 = df_Mt_p_data[df_Mt_p_data["Percentile"] <= 50]
    df_pdif_MtPlt = pd.concat([below50, above50])

    # Calculate %median bias
    medianBias = round(np.median(df_pdif_MtPlt["Percent Difference"]), 1)

    # Fit the mountain plot data with GAM
    x_fit = df_pdif_MtPlt["Percent Difference"]
    y_fit = df_pdif_MtPlt["Percentile"]
    gam = LinearGAM(s(0, n_splines=n_splines, spline_order=spline_order)).fit(
        x_fit, y_fit
    )

    XX = gam.generate_X_grid(term=0, n=5000)

    # Create the Plotly figure
    fig = px.scatter(
        df_pdif_MtPlt,
        x="Percent Difference",
        y="Percentile",
        color="Color_var",
        color_continuous_scale="Bluered_r",
        labels={"Y-X": "%Difference", "Percentile": "Percentile (%)"},
        title="Mountain Plot",
    )

    # Add median bias line
    fig.add_shape(
        type="line",
        x0=medianBias,
        y0=0,
        x1=medianBias,
        y1=50,
        line=dict(color="red", dash="solid", width=1),
        name="Median Bias",
    )

    # Add reference line
    fig.add_shape(
        type="line",
        x0=0,
        y0=0,
        x1=0,
        y1=50,
        line=dict(color="black", dash="dash", width=0.75),
        name="Reference",
    )

    # Annotate median bias
    fig.add_annotation(
        x=medianBias,
        y=52,
        text=f"Median Bias: {medianBias}",
        showarrow=False,
        font=dict(color="red"),
    )

    # Add Penalized Cubic Regression Splines prediction line
    fig.add_trace(
        go.Scatter(
            x=XX.flatten(),
            y=gam.predict(XX),
            mode="lines",
            line=dict(color="grey", dash="solid", width=1),
            name="Penalized Cubic Splines Prediction",
        )
    )

    fig.update_layout(
        legend=dict(
            title="Color_var",
            orientation="h",
            yanchor="bottom",
            y=1.02,
            xanchor="right",
            x=1,
        )
    )

    return fig, df_pdif_MtPlt


"""
Example usage of the function:
fig, df_pdif_MtPlt = mountain_plot_pd2_cubic_spline(x = df_MC['TBIL'],
                                       y = df_MC['TBIL2'],
                    z = df_MC['TBIL'],
                    n_splines=10,
                    spline_order=3)
fig.show()
print(df_pdif_MtPlt.head())


"""


def mountain_plot_ad2_cubic_spline(x, y, z=None, n_splines=12, spline_order=3):
    """
    Generate a Plotly mountain plot for simple differences using Penalized Cubic Regression Splines.

    Args:
        x (array-like): Reference method data.
        y (array-like): Test method data.
        z (array-like, optional): Category labels for coloring. If None, uses x.

    Returns:
        tuple: (plotly.graph_objects.Figure, pd.DataFrame)
            The Plotly figure object and the DataFrame containing the plot data.
    """

    if z is None:
        z = x

    # Create a new DataFrame
    df_Mt = pd.DataFrame(
        {x.name: x, "Test Method (Y)": y, "Color_var": z}
    )

    # Calculate differences and percent differences
    df_Mt["Y-X"] = df_Mt["Test Method (Y)"] - df_Mt[x.name]

    ## Generate DataFrame for Difference plot
    df_Mt_p_data = df_Mt.sort_values(by="Y-X")
    df_Mt_p_data["Rank"] = df_Mt_p_data["Y-X"].rank(ascending=True)
    df_Mt_p_data["Percentile"] = df_Mt_p_data["Rank"] * 100 / (len(df_Mt) + 1)

    # Split into above and below 50th percentile
    above50 = df_Mt_p_data[df_Mt_p_data["Percentile"] > 50].copy()
    above50["Percentile"] = 100 - above50["Percentile"]
    below50 = df_Mt_p_data[df_Mt_p_data["Percentile"] <= 50]
    df_pdif_MtPlt = pd.concat([below50, above50])

    # Calculate median bias
    medianBias = round(np.median(df_pdif_MtPlt["Y-X"]), 1)

    # Fit the mountain plot data with Penalized Cubic Regression Splines
    x_fit = df_pdif_MtPlt["Y-X"]
    y_fit = df_pdif_MtPlt["Percentile"]
    gam = LinearGAM(s(0, n_splines=n_splines, spline_order=spline_order)).fit(
        x_fit, y_fit
    )

    XX = gam.generate_X_grid(term=0, n=5000)

    # Create the Plotly figure
    fig = px.scatter(
        df_pdif_MtPlt,
        x="Y-X",
        y="Percentile",
        color="Color_var",
        color_continuous_scale="Bluered_r",
        labels={"Y-X": "Method Difference", "Percentile": "Percentile (%)"},
        title=""
    )

    # Add median bias line
    fig.add_shape(
        type="line",
        x0=medianBias,
        y0=0,
        x1=medianBias,
        y1=50,
        line=dict(color="red", dash="solid", width=1),
        name="Median Bias",
    )

    # Add reference line
    fig.add_shape(
        type="line",
        x0=0,
        y0=0,
        x1=0,
        y1=50,
        line=dict(color="black", dash="dot", width=0.75),
        name="Reference",
    )

    # Annotate median bias
    fig.add_annotation(
        x=medianBias,
        y=52,
        text=f"Median Bias: {medianBias}",
        showarrow=False,
        font=dict(color="red"),
    )

    # Add Penalized Cubic Regression Splines prediction line
    fig.add_trace(
        go.Scatter(
            x=XX.flatten(),
            y=gam.predict(XX),
            mode="lines",
            line=dict(color="grey", dash="solid", width=1),
            name="Penalized Cubic Splines Prediction",
        )
    )

    fig.update_layout(
        legend=dict(
            title="Color_var",
            orientation="h",
            yanchor="bottom",
            y=1.02,
            xanchor="right",
            x=1,
        )
    )

    return fig, df_pdif_MtPlt

def MC_output(document,analyte, x, y, z=None, z2=None, Unit=None, **error_info):    
    """
    This function is for creatining the entire Method Comparison section;
    It outputs a Table summary, comparing the difference and %difference given the error conditions
    It outputs a Summary and Table of parired Student T-test
    It outputs Histrogram and Q-Q plots for distribution analysis of each axis
    it outputs a Regression Analysis by Deming, a Difference plot, and %Difference plot
    """

    error1 = error_info.get("error1", None)
    error2 = error_info.get("error2", None)
    Error_level_cut_off = error_info.get("Error_level_cut_off", None)

    # adjust page margins
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(0.5)
        section.right_margin = Cm(0.5)
        
    # add heading
    #p = document.add_heading(f"Method Comparison Studies", level=1)
    p = document.add_heading(f" {y.name} vs. {x.name}", level=2)
    p = document.add_paragraph("")

    p.add_run(
        "Pass/Fail is determined based on the allowable error, which is defined as the following."
    )
    p = document.add_paragraph("")
    if Error_level_cut_off == 0 or Error_level_cut_off == None:
        p.add_run("Across all range of the measuring interval: ")
        if error2 is not None:
            p.add_run(str(error2))
            p.add_run("%")
        elif error1 is not None:
            p.add_run(" ± ")
            p.add_run(str(error1))
            p.add_run(" ")
            p.add_run(Unit)
    else:
        p.add_run("< ")
        p.add_run(str(Error_level_cut_off))
        p.add_run(": ± ")
        p.add_run(str(error1))
        p.add_run(" ")
        p.add_run(Unit)

        p = document.add_paragraph("")
        p.add_run("≥ ")
        p.add_run(str(Error_level_cut_off))
        p.add_run(": ±")
        p.add_run(str(error2))
        p.add_run("%")

    p = document.add_paragraph("")
    p.add_run("The Reference Method (X) is: ")
    p.add_run(str(x.name)).bold = True

    p = document.add_paragraph("")
    p.add_run("The Test Method (Y) is: ")
    p.add_run(str(y.name)).bold = True

    ### Distribution Analysis
    #### Histogram
    #document.add_page_break()
    document.add_heading("Distribution of Measurements:", level=3)
    Histogram_grouped(x, y)
    FigWhiteCrop("Figures/boxplot.png")
    document.add_picture("Figures/cropped_plot_boxplot.png", width=Inches(8))

    Table1, Table1b = sample_comparions_tableX(x, y, z, **error_info)

    # Table summary of Distribution Analysis
    Table1b_summary = data_summary(Table1b)
    create_word_table(document, Table1b_summary)


    """
    Here we have t-test done if the Shapiro-Wilk test indicates the 
    measures by reference method is Guass
    """
    #####
    # Shapiro-Wilk Test on the difference between the Reference and Test method.
    Table1 = Table1.dropna(subset=[x.name, y.name])
    stat, p = shapiro(Table1[y.name] - Table1[x.name])
    print(
        f"Shapiro-Wilk Test Statistic on Test Method (Y) -x.name {stat}, p-value: {p}"
    )
    # intepretation of Shapiro-Wilk test
    alpha = 0.05
    if p > alpha:
        ####t-test
        #document.add_page_break()
        p = document.add_paragraph(
            "The difference in measurements between the Test Method and Reference Method sufficiently resembles a Gaussian distribution, based on Shapiro-Wilk Test. Consider the Student T-test."
        )
        document.add_heading("Summary of Paired Student t-Test: ", level=4)
        p = document.add_paragraph(
            "Null Hypothesis, H0 = The means of the methods are equivalent"
        )
        p = document.add_paragraph(
            "Alternate Hypothesis, H1 = The means of the methods are not equivalent"
        )

        # prepare data for t-test
        table_tt = Table1[[x.name, y.name]].dropna() # make sure no NaN values
        table_tt[x.name] = pd.to_numeric(table_tt[x.name], errors='coerce') # convert to numeric, coerce errors to NaN
        table_tt[y.name] = pd.to_numeric(table_tt[y.name], errors='coerce') # convert to numeric, coerce errors to NaN
        table_tt = table_tt[[x.name, y.name]].dropna() # drop NaN values again after conversion
        # from pingouin library, applying t-test
        t_Test_Results = ttest(
            table_tt[x.name], table_tt[y.name], paired=True
        ).round(4)
        t_Test_Results
        create_word_table(document, t_Test_Results)

        p = document.add_paragraph("")
        p = document.add_paragraph("T = t-value;  ")
        p.add_run("dof =  degrees of freedom:  ")
        p.add_run("p-val = p-value:  ")
        p.add_run("alternative = alternative of the test;  ")
        p.add_run("CI95% = confidence intervals of the difference in means;  ")
        p.add_run("cohen-d = Cohen d effect size;  ")
        p.add_run("BF10 = Bayes Factor of the alternative hypothesis;  ")
        p.add_run("power = achieved power of the test ( = 1 - type II error);  ")

        p_val = t_Test_Results['p_val'].values[0]
        if p_val > 0.05:
            p = document.add_paragraph()
            p.add_run("Accept ").bold = True
            p.add_run(
                "the Null Hypothesis, that there are no significant differences between the means of the methods (p>0.05)."
            )

            p = document.add_paragraph()
            p.add_run("Note: ").bold = True
            p.add_run(
                "The t-test result is valid only when there is sufficient amount of normality across the distribution of the measurements. "
            )
            p.add_run(
                "Consider the Q-Q plots, histrogram distribution and other statistical procedures before accepting the conclusions above."
            )
        else:
            p = document.add_paragraph()
            p.add_run("Reject ").bold = True
            p.add_run(
                "the Null Hypothesis, that there are no significant differences between the means of the methods (p<0.05)."
            )
            p = document.add_paragraph(
                "ACCEPT the Alternative Hypothesis, that there are significant differences between the means of the methods."
            )

            p = document.add_paragraph()
            p.add_run("Note: ").bold = True
            p.add_run(
                "The t-test result is valid only when there is sufficient amount of normality across the distribution of the measurements. "
            )
            p.add_run(
                "Consider the Q-Q plots, histogram distribution and other statistical procedures before accepting the conclusions above."
            )
    else:
        document.add_heading("Shapiro-Wilk Test: ", level=3)
        p = document.add_paragraph(
            "The measurements of the difference between the Test Method and Reference Method do NOT sufficiently resemble a Gaussian distribution, as based on Shapiro-Wilk Test. "
        )
        p.add_run("Proportional biases may be present.")
    
    ### Regression and Bias Analysis
    #document.add_page_break()
    document.add_heading("Regression & Bias Analysis", level=3)
    # Deming Regression assuming equal variance plot
    n, slope, yintercept, Pearson_r = Deming_Plot_Equal_Variance_with_Error2_PX(x, y, z, z2, **error_info)

    p = document.add_paragraph("")
    p.add_run("Regression is based on Deming method, assuming equal variance of both methods. ")
    document.add_picture("Figures/cropped_plot_Deming_regression.png", width=Inches(6.5))


    
    # Summary Table of the Methods
    p = document.add_paragraph("")
    p.add_run("Comparison of Paired Measurements").bold = True
    create_word_table(document, Table1)
    document.add_page_break()

    # Difference Plot
    median_bias = Difference_plot_median_with_error2(x, y, z, z2, **error_info)
    # %Difference Plot
    percent_bias = Difference_plot_Percent_median_with_error2(x, y, z, z2, **error_info)


    p = document.add_paragraph("")
    p.add_run("Difference Plots").bold = True

    table = document.add_table(rows=2, cols=1)

    cell_1 = table.cell(0, 0)
    p = cell_1.paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run = p.add_run()
    run.add_picture("Figures/cropped_plot_Difference_Plot_median.png", width=Inches(6))
    # Add a line break
    p = cell_1.add_paragraph("")

    cell_2 = table.cell(1, 0)
    p = cell_2.paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run()
    run.add_picture(
        "Figures/cropped_plot_Difference_Plot_median_per.png", width=Inches(6)
    )
    


    write_to_csv(analyte, n, round(slope,4), round(yintercept,4), round(Pearson_r,4), round(median_bias,4), round(percent_bias,4), Unit)
    document.save("MethodComparison.docx")

def precision_output(document, analyte, df, Unit):
    """
    Function to output precision data to the document.
    The dataframe should be in long format with columns 'Day', 'Run', 'Replicate', and 'Result'.
    Outputs the ANOVA table. 
    Outputs the derived within run, between run, and between day standard deviations and %CVs.
    Outputs the total standard deviation and %CV.
    Outputs a LJ chart, with SD that is based on grand SD.
    """
    

    model = ols('Result ~ C(Day) / C(Run)', data=df).fit()
    #model = ols('Result ~ C(Day) + C(Day)/C(Run)', data=df).fit() 
    anova_table = sm.stats.anova_lm(model, type=2) 
    #print(anova_table.to_string())
    
    # Extract mean squares from the ANOVA table
    MS_error = anova_table.loc['Residual','mean_sq']
    MS_run = anova_table.loc['C(Day):C(Run)','mean_sq']
    MS_day = anova_table.loc['C(Day)','mean_sq']

    # Within run variation
    SD_Repeatability = (MS_error)**0.5
    SD_Repeatability

    #get the mean of the 'Result' column
    meanLX = df['Result'].mean()
    # get number replicates and runs
    n_replicate = max(df['Replicate'])
    n_run = max(df['Run'])
    n = len(df)
    
    # Within Run CV
    wi_runCV = 100 * (SD_Repeatability / meanLX)
    wi_runCV

    # Between run variation
    SD_Run = (MS_run - MS_error) / n_replicate
    if SD_Run < 0:
            SD_Run = 0
            bw_runCV = 0
    else:
        SD_Run = SD_Run**0.5
         # Between Run CV
        bw_runCV = 100 * (SD_Run / meanLX)
    SD_Run

   

    # Between Day variation
    SD_Day = (MS_day - MS_run) / (n_replicate * n_run)
    if SD_Day < 0:
        SD_Day = 0
        bw_dayCV = 0
    else:
        SD_Day = SD_Day**0.5
        # Between day CV
        bw_dayCV = 100 * (SD_Day / meanLX)
        bw_dayCV

        
    
    # Total Variation as SD
    SD_total = ((SD_Repeatability**2) + (SD_Run**2) + (SD_Day**2)) ** 0.5
    # Total %CV
    Total_CV = ((wi_runCV**2) + (bw_runCV**2) + (bw_dayCV**2)) ** (0.5)

    # Print ANOVA output
    anova_table_df = pd.DataFrame({
        'Source of Variation': ['Day', 'Run(Day)', 'Error'],
        'Degrees of Freedom (DF)': anova_table['df'],
        'Sum of Squares (SS)': anova_table['sum_sq'].round(4),
        'Mean Squares (MS)': anova_table['mean_sq'].round(5),
        'F Value': [f"{x:.3f}" if not pd.isna(x) else "" for x in anova_table['F']],
        'P-value': [f"{x:.3e}" if not pd.isna(x) else "" for x in anova_table['PR(>F)']]
    })
    print(anova_table_df.to_string(index=False))

    # Create a DataFrame for the variation summary
    variation_df = pd.DataFrame({
        'Source of Variation': ['Within Run', 'Between Run', 'Between Day', 'Total'],
        'Standard Deviation (SD)': [SD_Repeatability, SD_Run, SD_Day, SD_total],
        'Coefficient of Variation (CV) %': [wi_runCV, bw_runCV, bw_dayCV, Total_CV]
    }).round(4)
    print(variation_df.to_string(index=False))

    # Write the precision assessment to the document
    document.add_page_break()
    heading_title = str("Precision Assessment at" + " " + str(round(meanLX, 3)) + " " + Unit)
    document.add_heading(heading_title, level=2)
    
    # provide the raw data as a table
    p = document.add_paragraph('')
    p.add_run('The table below is a summary of qualifying raw data, where data with only 2 runs and 2 replicates per run are available.')
    create_word_table(document, df)
    
    ###############
    # Add LJ plots    
    # Ensure Replicate column is a cateogorical before plotting
    if 'Replicate' in df.columns:
        QC_replicate = df['Replicate'].astype(str)  
    else:
        QC_replicate = None
    # Ensure Run column is a cateogorical before plotting
    if 'Run' in df.columns:
        QC_run = df['Run'].astype(str)  
    else:
        QC_run = None
    
    #Add section heading to word document
    p.add_run(" ").bold = True
    p = document.add_paragraph('')
    p = document.add_paragraph('')
    p.add_run('Levey-Jennings Plot').bold = True
    #plot data on LJ chart and write to document
    LJ_plot2(
        df=df,
        x=df['Day'],
        y=df['Result'],
        Replicate=QC_replicate,
        QCRun=QC_run,
    )
    
    #document.add_picture('Figures/cropped_plot_LJ-Plot.png', width=Inches(6.5))
    # Add a paragraph for the image and center it
    image_paragraph = document.add_paragraph()
    run = image_paragraph.add_run()
    run.add_picture('Figures/cropped_plot_LJ-Plot.png', width=Inches(6.5))
    image_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    ###############

    ########## Add ANOVA table ##############
    p = document.add_paragraph('ANOVA Table:').bold = True
    create_word_table(document, anova_table_df)
    p = document.add_paragraph('')
    #p.add_run(anova_table_df.to_string(index=False))
    
    # Add to the document the variation df
    p = document.add_paragraph('')
    p.add_run('Variation Summary:').bold = True
    p = document.add_paragraph()
    p.add_run("At a mean of ")
    p.add_run(str(round(meanLX,3))).bold = True
    p.add_run(" ").bold = True
    p.add_run(Unit).bold = True
    p.add_run(", the following precision results are obtained.")
    create_word_table(document, variation_df)
    
    p = document.add_paragraph('')
    p = document.add_paragraph('')
    p.add_run("The total within laboratory imprecision at a mean of ")
    p.add_run(str(round(meanLX, 4))).bold = True
    p.add_run(" ± ")
    p.add_run(str(variation_df.iloc[3,1]))
    p.add_run(" ")
    p.add_run(Unit)
    p.add_run(", has a CV of ")
    p.add_run(str(variation_df.iloc[3,2])).bold = True
    p.add_run("%.")
    
    #Write the results summar to a CSV file
    # Define the header for the CSV file
    Total_CV = variation_df.iloc[3,2]
    Total_SD = variation_df.iloc[3,1]
    
    def write_to_csv_precision(analyte, n, meanLX, Total_CV, Total_SD, Unit):
        header = ['Analyte', 'N', 'Mean', 'Total %CV', 'Total SD' 'Unit']
    
        # Check if file exists to determine if we need to write the header
        filename = "PrecisionSummary.csv"
        file_exists = os.path.isfile(filename)
    
        # Open the file in append mode
        with open(filename, mode='a', newline='') as file:
            writer = csv.writer(file)
            
            # Write header if file doesn't exist
            if not file_exists:
                writer.writerow(header)
                
            # Write the data row
            writer.writerow([analyte, n, meanLX, Total_CV, Total_SD, Unit])
            
    write_to_csv_precision(analyte, n, round(meanLX,4), round(Total_CV,4), round(Total_SD,4), Unit)
    
    document.save("MethodComparison.docx")

def linearity_output(document, analyte, Measured_mean,Expected_mean, Measured_err, n,Unit, **error_info ):
    document.add_page_break()
    document.add_heading("Linearity Analysis", level=1)
    # p = document.add_paragraph('')
    p = document.add_paragraph('Regression is based on Ordinary Least Squares (OLS) regression. Shaded region represents the defined allowable performance limits.')
    
    linearity_plots(Measured_mean,Expected_mean, Measured_err, n, **error_info)
    document.add_picture('Figures/cropped_plot_Linearity.png', width=Inches(8))   

    
    df_Linearity_summary = pd.DataFrame({'Measured Mean': Measured_mean,
                                        'Expected Mean': Expected_mean, 
                                        'Measured Mean Error': Measured_err
                                        })
    
    df_Linearity_summary = df_Linearity_summary.round(3)
    p = document.add_paragraph('Summary of Linearity Results')
    create_word_table(document, df_Linearity_summary)
    p = document.add_paragraph('')
    p = document.add_paragraph('')
    p.add_run('Linear interval assessed:')
    linear_interval = f' {round(Measured_mean.min(),2)} - {round(Measured_mean.max(),2)} {Unit}' 
    p.add_run(linear_interval).bold = True
    p = document.add_paragraph('')
    p.add_run('Linear interval assessed with error:')
    linear_interval_err = f' {round((Measured_mean - Measured_err).min(),2)} - {round((Measured_mean + Measured_err).max(),2)} {Unit}'
    p.add_run( linear_interval_err).bold = True
    p = document.add_paragraph('')
    
    def write_to_csv_linearity(analyte, linear_interval, linear_interval_err, Unit):
        header = ['Analyte', 'Linear Interval', 'Linear Interval with Error', 'Unit']
    
        # Check if file exists to determine if we need to write the header
        filename = "LinearitySummary.csv"
        file_exists = os.path.isfile(filename)
    
        # Open the file in append mode
        with open(filename, mode='a', newline='') as file:
            writer = csv.writer(file)
            
            # Write header if file doesn't exist
            if not file_exists:
                writer.writerow(header)
                
            # Write the data row
            writer.writerow([analyte, linear_interval, linear_interval_err, Unit])
    write_to_csv_linearity(analyte, linear_interval, linear_interval_err, Unit)
    document.save("MethodComparison.docx")

def RI_output(document, analyte, df, low_lim_RI, high_lim_RI, Unit):   
    '''
    _summary_
        This function outputs a page on reference interval verification evlauation.
        
        Must have a data frame with "Result" and "Gender" columns
        Gender must have values as "Male" and/or "Female"

    ''' 
    document.add_page_break()
    document.add_heading(f"Reference Interval Verification of {analyte}", level=1)
    # p = document.add_paragraph('')
    p = document.add_paragraph('Reference interval (RI) verification is as recommended by CLSI EP28-A3c. ')
    p.add_run('Ideally at least 90% of the results must be within the proposed RI.')

    p = document.add_paragraph("")
    p.add_run(f"Proposed RI: ")
    p.add_run(f"{low_lim_RI} - {high_lim_RI} {Unit}").bold = True
    
    df['Within RI'] = df['Result'].apply(lambda val: 'Yes' if low_lim_RI <= val <= high_lim_RI else 'No')

    # CLSI EP28-A3c guidelines calls 2 out of 20 as allowable maximum outliers
        # Which is essential 90% of the results need to be inside the RI
    # determine what precentage is outside of RI Limits
    inlier = len(df[df["Within RI"]=='Yes'])
    total_count = len(df)
    pWithin = round(100*inlier/total_count,1)
    print(f"inlier = {inlier}; + total RI count = {total_count}; Precentage inside = {pWithin}")
    
    create_word_table(document,df)
    
    p = document.add_paragraph("")
    p = document.add_paragraph("")
    # write if the proposed RI meets acceptable criteria.
    if pWithin >= 90:
        p.add_run("The proposed RI may be suitable, given that ")
        p.add_run(str(pWithin)).bold = True
        p.add_run("%").bold = True
        p.add_run(" of the reference individuals are")
        p.add_run(" within the proposed RI.").bold = True
    else:
        p.add_run(f"The proposed reference interval may not be ideal(inlier/total = {inlier}/{total_count} = {pWithin:.1f}%).  Additional sampling and/or optimization may need to be considered.")
    
    medians_df = RI_histogram(df, low_lim_RI, high_lim_RI)
    print(medians_df)
    p = document.add_paragraph("")
    if medians_df.shape[0] == 2:
        if (medians_df.iloc[0,0] == "Female") and (medians_df.iloc[1,0] == "Male"): 
            p.add_run(f"The median value by gender is {round(medians_df.iloc[0,1],1)} {Unit} and {round(medians_df.iloc[1,1],1)} {Unit}, for {medians_df.iloc[0,0]} and {medians_df.iloc[1,0]}, respectively.")
    else:
        p.add_run(f"The median value of the reference individuals is {round(medians_df.iloc[0,1],1)} {Unit}.")
   # CC.create_word_table(document, medians_df)
    
    document.add_picture("Figures/RI.png", width=Inches(7.5))

    document.save("MethodComparison.docx")
    

### plotting and fitting an exponential function
"""
This function helps to streamline plotting and fitting an exponential function.
For example, this can be useful when it comes to 
plotting CV as a function of concentrion. i.e., a precision profile curve

"""


def precision_LoQ_linefit(x, y, x_valuesFit):
    # take the natural log of the data
    logx = np.log(x)
    logy = np.log(y)

    # determine the lienar regression fit for x and y data; already natural log transformed
    reg_summary = stats.linregress(logx, logy)
    print(f"Fitted Line: y = {reg_summary.slope}x+{reg_summary.intercept}")
    print(f"R-squared: {reg_summary.rvalue**2}")

    plt.figure(figsize=(8, 6))
    sns.scatterplot(x=logx, y=logy, label="Data Points")
    sns.lineplot(
        x=logx,
        y=reg_summary.slope * logx + reg_summary.intercept,
        color="red",
        label=f"Fitted Function:  ln y = {reg_summary.slope:.5f}ln x+{reg_summary.intercept:.5f} \n R$^2$: {reg_summary.rvalue**2:.5f}",
    )


def precision_LoQ_powerFit(x, y, x_valuesFit):
    # take the natural log of the data
    logx = np.log(x)
    logy = np.log(y)

    from scipy import stats

    reg_summary = stats.linregress(logx, logy)

    # take the natural log of the data
    logx = np.log(x)
    logy = np.log(y)
    a = np.e**reg_summary.intercept
    b = reg_summary.slope
    sns.scatterplot(x=x, y=y, label="Data Points")
    # sns.regplot(x = x, y=y, label='log transformed data')
    sns.lineplot(
        x=x_valuesFit,
        y=a * (x_valuesFit ** (b)),
        color="red",
        label=f"Fitted Function:  $y = {a}x^{{{b}}}$ \n R$^2$: {reg_summary.rvalue**2:.5f}",
    )
    # logy = reg_summary.intercept + reg_summary.slope*logx
    functional_sensitivity = (20 / a) ** (1 / b)
    print(
        f"The functional sensitivity, or the average level at 20% CV, is approximately {functional_sensitivity}"
    )
    return a, b, functional_sensitivity


def precision_LoQ_powerFit_Output(document, x, y, x_valuesFit, unit, ylimit):
    # add heading
    p = document.add_heading("Limit of Quantitation (LoQ) Studies", level=2)

    MyTable = pd.DataFrame({"Average Measurement": x, "%CV": y})
    create_word_table(document, MyTable)
    p = document.add_paragraph("")

    document.add_heading("Precision Profile", level=3)

    # take the natural log of the data
    logx = np.log(x)
    logy = np.log(y)

    from scipy import stats

    reg_summary = stats.linregress(logx, logy)

    # take the natural log of the data
    logx = np.log(x)
    logy = np.log(y)
    a = np.e**reg_summary.intercept
    b = reg_summary.slope

    plt.figure(figsize=(6, 6))
    sns.set_style("whitegrid")
    sns.scatterplot(x=x, y=y, label="Data Points")  # plot the data
    # plot the fitted line
    sns.lineplot(
        x=x_valuesFit,
        y=a * (x_valuesFit ** (b)),
        color="red",
        label=f"Fitted Function:  $y = {a:.4f}x^{{{b:.4f}}}$ \n R$^2$: {reg_summary.rvalue**2:.4f}",
    )
    plt.ylim(ylimit)
    plt.savefig("LoQ.png", facecolor="white", format="png")
    plt.close()

    document.add_picture("LoQ.png", width=Inches(6))

    # calculate the functional sensitivity
    functional_sensitivity = (20 / a) ** (1 / b)

    p = document.add_paragraph("")
    p.add_run(
        "The functional sensitivity, which is the average expected measurement at 20% CV, is approximately "
    )
    p.add_run(str(round(functional_sensitivity, 4))).bold = True
    p.add_run(" ").bold = True
    p.add_run(unit)
    p.add_run(".")

    document.save("MethodComparison.docx")

    return a, b, round(functional_sensitivity, 4)


### This function creates a Word document with the signature and approval section
### It uses the python-docx library

def page_singature_approval(Test_Analyte,
                            Test_Reagent_Manufacturer,
                            Test_Instrument,
                            Test_Instrument_sn,
                            Reference_Analyte,
                            Reference_Reagent_Manufacturer,
                            Reference_Instrument,
                            Hospital_info):
    
    
    from docx import Document
    from docx.shared import Inches, Cm, Pt

    #open a blank document
    document = Document()

    #adjust page margins
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    #add heading
    p = document.add_heading('Verification of ', level=0)
    p.add_run(Test_Analyte)


    document.add_heading('Summary of Approval and Signature', level = 1)
    p = document.add_paragraph('')
    p = document.add_paragraph('')

    p.add_run('This document is prepared for')
    p = document.add_paragraph('')
    p.add_run('    Lakeridge Health')
    p = document.add_paragraph('')
    p.add_run('    Ajax-Pickering Hospital')
    p = document.add_paragraph('')
    p.add_run('    Laboratory Medicine & Pathology')

    #This to type the following text:
    ################
    '''
    This is a verification of "Test Analyte" from "Reagent Manufacturer" on the "Instrument" at the "Hospital Site". 
    The deatails of the verification study scheme(s) and plans are detailed separately and/or are performed according to the advice of the attending clinical biochemist.
    The key studies here are the Method Comparison for relative accuracy assesment or method agreement, 
    as well as precision studies to evaluate repeatability and/or reproducibility.
    The method comparison involves evaluation of the test method against the results obtained from the reference method, 
    which in this case is "Reference Test" from "Reference Reagent Manufacturer" on the "Reference Instrument".
    '''
    #################
    
    

    p = document.add_paragraph('')
    p = document.add_paragraph('')
    p.add_run('This is a verification of ')
    p.add_run(Test_Analyte).bold = True
    p.add_run(' from ')
    p.add_run(Test_Reagent_Manufacturer).bold = True
    p.add_run(' on the ')
    p.add_run(Test_Instrument).bold = True
    p.add_run(' (Instrument s/n: ')
    p.add_run(Test_Instrument_sn).bold = True
    p.add_run(' )')
    p.add_run(' at ')
    p.add_run(Hospital_info)
    p.add_run('. ')
    
    messaage ="""The details of the verification study scheme(s) and plans are noted separately and/or are performed according to the advice of the attending clinical biochemist. The key studies are Method Comparison for relative accuracy assessment and/or method agreement, as well as precision studies to evaluate repeatability and/or reproducibility. """
    #p.add_run('The details of the verification study scheme(s) and plans are noted separately and/or are performed according to the advice of the attending clinical biochemist. ')
    #p.add_run('The key studies are Method Comparison for relative accuracy assesment  and/or method agreement, as well as precision studies to evaluate repeatability and/or reproduceability. ')
    p.add_run(messaage)
    p.add_run('The method comparison involves evaluation of the test method against the results obtained from the reference method, which in this case is ')
    p.add_run(Reference_Analyte)

    p.add_run(' from ')
    p.add_run(Reference_Reagent_Manufacturer)
    p.add_run(' on the ')
    p.add_run(Reference_Instrument)
    p.add_run('.')



    ##This to type the following text:
    '''
    The data presented below has been reviewed and analyzed to the best of my knowledge.
    Some major changes, if any, are noted below.
        * None
    '''
    p = document.add_paragraph('')
    p.add_run('The data presented has been reviewed and analyzed to the best of my knowledge. Some major changes, if any, are noted below. ')
    p = document.add_paragraph('None', style = 'List Bullet' )
    p.paragraph_format.left_indent = Pt(36) # indent the bulleted point

    p = document.add_paragraph('')
    p.add_run('Withstanding any unforeseen concerns, and based on the studies presented, analyzed and reviewed here, I have approved the test as acceptable for clinical use. ')
    p = document.add_paragraph('')
    p = document.add_paragraph('')
    p = document.add_paragraph('')


    ######Signature ######
    p.add_run('Rajeevan Selvaratnam').bold = True
    p.add_run(", PhD, NRCC, DABCC, FCACB, FADLM").font.size = Pt(8)
    p=document.add_paragraph('')
    p.add_run('Clinical Biochemist - Laboratory Medicine Program, University Health Network').italic = True
    p=document.add_paragraph('')
    p.add_run('Assistant Professor - Laboratory Medicine& Pathobiology, University of Toronto').italic = True
    p=document.add_paragraph('')
    p.add_run('Date: ').bold = True
    document.add_page_break()

    document.save('MethodComparison.docx')
    
    return Test_Analyte
